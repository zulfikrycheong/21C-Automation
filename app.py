import base64
from datetime import datetime
import io
import json
import os
import re
import sqlite3
import repo_sync

import docx
from google.oauth2.service_account import Credentials
import gspread
import pypdfium2 as pdfium
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
import streamlit as st
import streamlit.components.v1 as components

# --- APP CONFIGURATION & BRANDING ---
st.set_page_config(
    page_title="ChambersOS | Practice Operations Suite",
    page_icon="⚖️",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# Custom Executive Theme Injection (Theme-Aware & Native Streamlit Compatible)
st.markdown(
    """
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');
        
        /* Global Typography */
        html, body, [class*="css"] {
            font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif;
        }
        
        .main-header {
            padding: 0.8rem 0rem 1.4rem 0rem;
            border-bottom: 1px solid rgba(128, 128, 128, 0.2);
            margin-bottom: 1.5rem;
        }
        
        .suite-title {
            font-size: 1.85rem;
            font-weight: 800;
            letter-spacing: -0.5px;
            margin: 0;
        }
        
        .suite-subtitle {
            font-size: 0.95rem;
            font-weight: 500;
            margin-top: 4px;
        }
        
        /* --- EXECUTIVE WING TABS (Clean Brushed Gold Overhaul) --- */
        .stTabs [data-baseweb="tab-list"] {
            gap: 12px;
            border-bottom: 2px solid rgba(128, 128, 128, 0.2);
            padding-bottom: 2px;
            overflow: visible !important;
            background-color: transparent !important;
        }
        
        .stTabs [data-baseweb="tab"] {
            height: 48px;
            white-space: nowrap !important;
            border-radius: 6px 6px 0px 0px;
            padding: 10px 22px !important;
            font-weight: 600;
            font-size: 0.92rem;
            border: none !important;
            background-color: rgba(128, 128, 128, 0.08) !important;
            overflow: visible !important;
            outline: none !important;
            box-shadow: none !important;
            transition: all 0.2s ease;
        }
        
        .stTabs [data-baseweb="tab"]:hover {
            color: #d4af37 !important;
        }
        
        .stTabs [data-baseweb="tab"]:focus, 
        .stTabs [data-baseweb="tab"]:active {
            outline: none !important;
            box-shadow: none !important;
        }
        
        .stTabs [data-baseweb="tab"] div {
            overflow: visible !important;
            color: inherit !important;
        }
        
        /* Active Wing Tab: Brushed Gold Bottom Accent */
        .stTabs [aria-selected="true"] {
            color: #d4af37 !important;
            border-bottom: 3px solid #d4af37 !important;
            background-color: rgba(212, 175, 55, 0.08) !important;
            outline: none !important;
            box-shadow: none !important;
        }
        
        /* Kill Native Streamlit Red/Blue Focus Rings & Replace with Brushed Gold */
        input:focus, select:focus, textarea:focus, button:focus {
            border-color: #d4af37 !important;
            box-shadow: 0 0 0 1px #d4af37 !important;
        }
        
        /* Fix Metric Value Text */
        [data-testid="stMetricValue"] {
            font-size: 1.65rem !important;
            white-space: normal !important;
            word-break: break-word !important;
        }
        
        /* Upload Area Styling */
        [data-testid="stFileUploaderDropzone"] {
            padding: 3.5rem 2rem !important;
            border-radius: 10px !important;
            border: 2px dashed rgba(128, 128, 128, 0.3) !important;
            transition: all 0.2s ease-in-out;
        }
        [data-testid="stFileUploaderDropzone"]:hover {
            border-color: #d4af37 !important;
        }
    </style>
    """,
    unsafe_allow_html=True,
)

# --- WORKSPACE TOP BAR ---
c_logo, c_title = st.columns([1, 4])
with c_logo:
    if os.path.exists("Company Logo.png"):
        st.image("Company Logo.png", width=220)
with c_title:
    st.markdown(
        """
        <div class="main-header">
            <h1 class="suite-title">⚖️ ChambersOS Legal Automation Hub</h1>
            <div class="suite-subtitle">Integrated Matter Intake Engine & Crown Worldwide Archival Terminal</div>
        </div>
        """,
        unsafe_allow_html=True,
    )

# --- 1. GOOGLE SHEETS MATRIX SYNC ---
GOOGLE_SHEET_NAME = "21 Chambers LLC Client List"
SHEET_TAB_NAME = datetime.now().strftime("%B %Y")
CURRENT_YEAR = datetime.now().strftime("%Y")


def get_google_sheet():
    scopes = [
        "https://www.googleapis.com/auth/spreadsheets",
        "https://www.googleapis.com/auth/drive",
    ]
    if os.path.exists("credentials.json"):
        creds = Credentials.from_service_account_file(
            "credentials.json", scopes=scopes
        )
    else:
        encoded_str = st.secrets["encoded_creds"]
        decoded_bytes = base64.b64decode(encoded_str)
        creds_dict = json.loads(decoded_bytes)
        creds = Credentials.from_service_account_info(creds_dict, scopes=scopes)

    client = gspread.authorize(creds)
    workbook = client.open(GOOGLE_SHEET_NAME)

    try:
        sheet = workbook.worksheet(SHEET_TAB_NAME)
    except gspread.exceptions.WorksheetNotFound:
        template_sheet = workbook.worksheet("Template")
        sheet = workbook.duplicate_sheet(
            source_sheet_id=template_sheet.id,
            new_sheet_name=SHEET_TAB_NAME,
            insert_sheet_index=0,
        )
        headers = [[
            "Index",
            "Date Opened",
            "File / Matter No.",
            "Matter Type",
            "Client Name(s)",
            "Contact Details",
            "Referral Source",
            "Logged to Matrix",
            "Remarks",
        ]]
        sheet.update(range_name="A1:I1", values=headers)

    return sheet


def get_next_matter_number(sheet):
    try:
        vals = sheet.col_values(3)[1:]
        valid = [int(v.strip()) for v in vals if str(v).strip().isdigit()]
    except Exception:
        valid = []

    matter_nos = sheet.col_values(3)
    mx, mx_idx = -1, 1
    for idx, val in enumerate(matter_nos):
        if str(val).strip().isdigit() and int(str(val).strip()) > mx:
            mx = int(str(val).strip())
            mx_idx = idx + 1

    target_row = mx_idx + 1 if mx != -1 else 2
    base_max = mx if mx != -1 else (max(valid) if valid else 20260728)
    next_no = str(base_max + 1)
    next_idx = target_row - 1
    return next_no, target_row, next_idx


# --- 2. VECTOR PDF ENGINE ---
def generate_perfect_pdf(
    matter_no, clients_text, contacts_text, matter_type, date_opened
):
  buffer = io.BytesIO()
  doc = SimpleDocTemplate(
      buffer,
      pagesize=A4,
      leftMargin=21.24,
      rightMargin=21.24,
      topMargin=36.0,
      bottomMargin=36.0,
  )

  style_normal_20 = ParagraphStyle(
      "Norm20",
      fontName="Times-Roman",
      fontSize=20,
      leading=26,
      alignment=TA_LEFT,
  )
  style_bold_20 = ParagraphStyle(
      "Bold20",
      fontName="Times-Bold",
      fontSize=20,
      leading=26,
      alignment=TA_LEFT,
  )
  style_firm_title = ParagraphStyle(
      "FirmTitle",
      fontName="Times-Bold",
      fontSize=20,
      leading=26,
      alignment=TA_CENTER,
  )
  style_firm_body = ParagraphStyle(
      "FirmBody",
      fontName="Times-Roman",
      fontSize=20,
      leading=26,
      alignment=TA_CENTER,
  )

  # Bumped size from 109 to 125 (~15% increase) for easier cutting
  style_giant_foot = ParagraphStyle(
      "GiantFoot",
      fontName="Helvetica-Bold",
      fontSize=125,
      leading=132,
      alignment=TA_CENTER,
  )

  story = []

  client_lines = [
      line.strip() for line in clients_text.split("\n") if line.strip()
  ]
  contact_lines = [
      line.strip() for line in contacts_text.split("\n") if line.strip()
  ]

  party_elements = []
  if len(client_lines) > 0:
    party_elements.append(Paragraph(client_lines[0], style_normal_20))
  if len(contact_lines) > 0:
    party_elements.append(Paragraph(contact_lines[0], style_normal_20))

  party_elements.append(Spacer(1, 14))

  if len(client_lines) > 1:
    party_elements.append(Paragraph(client_lines[1], style_normal_20))
  if len(contact_lines) > 1:
    party_elements.append(Paragraph(contact_lines[1], style_normal_20))

  col_w_left = 1.333 * 72
  col_w_right = 6.346 * 72

  top_table = Table(
      [["", party_elements]], colWidths=[col_w_left, col_w_right]
  )
  top_table.setStyle(
      TableStyle([
          ("GRID", (0, 0), (-1, -1), 1.5, colors.black),
          ("VALIGN", (0, 0), (-1, -1), "TOP"),
          ("TOPPADDING", (0, 0), (-1, -1), 6),
          ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
          ("LEFTPADDING", (0, 0), (-1, -1), 7.2),
          ("RIGHTPADDING", (0, 0), (-1, -1), 6),
      ])
  )
  story.append(top_table)
  story.append(Spacer(1, 24))

  story.append(Paragraph("21 CHAMBERS LLC", style_firm_title))
  story.append(Spacer(1, 4))
  story.append(
      Paragraph(
          "2 HAVELOCK ROAD #06-17<br/>HAVELOCK 2<br/>SINGAPORE 059763",
          style_firm_body,
      )
  )
  story.append(Spacer(1, 4))
  story.append(
      Paragraph(
          "TEL: 6224 1848 &nbsp;&nbsp;&nbsp;&nbsp;&nbsp; FAX: 6223 3092",
          style_firm_body,
      )
  )
  story.append(Spacer(1, 40))

  full_matter_name = "Uncontested Divorce"
  if matter_type == "CD":
    full_matter_name = "Contested Divorce"
  elif matter_type == "Annulment":
    full_matter_name = "Annulment"
  elif matter_type == "Variation":
    full_matter_name = "Variation"
  elif matter_type == "Others":
    full_matter_name = "Others"

  file_block_text = (
      f"{matter_no}<br/>Opening date: {date_opened}<br/>Closure date:"
  )

  matrix_rows = [
      [
          Paragraph("SUBJECT<br/>MATTER", style_bold_20),
          Paragraph(full_matter_name, style_normal_20),
      ],
      [
          Paragraph("FILE", style_normal_20),
          Paragraph(file_block_text, style_normal_20),
      ],
      [
          Paragraph("Legal Fee", style_normal_20),
          Paragraph("CASH", style_normal_20),
      ],
      [Paragraph("Remarks", style_normal_20), Paragraph("", style_normal_20)],
  ]

  b_col1 = 1.596 * 72
  b_col2 = 6.083 * 72

  bottom_table = Table(matrix_rows, colWidths=[b_col1, b_col2])
  bottom_table.setStyle(
      TableStyle([
          ("GRID", (0, 0), (-1, -1), 1.5, colors.black),
          ("VALIGN", (0, 0), (-1, -1), "TOP"),
          ("TOPPADDING", (0, 0), (-1, -1), 6),
          ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
          ("LEFTPADDING", (0, 0), (-1, -1), 7.2),
          ("RIGHTPADDING", (0, 0), (-1, -1), 6),
      ])
  )

  # LOCKED LAYOUT BLOCK: Nests bottom table and giant number with exact 2pt (2px visual equivalent) spacing
  locked_block_table = Table(
      [[bottom_table], [Paragraph(matter_no, style_giant_foot)]],
      colWidths=[7.942 * 72],
  )
  locked_block_table.setStyle(
      TableStyle([
          ("VALIGN", (0, 0), (-1, -1), "TOP"),
          ("TOPPADDING", (0, 0), (0, 0), 0),
          ("BOTTOMPADDING", (0, 0), (0, 0), 2),  # Exactly 2pt gap below table
          ("TOPPADDING", (0, 1), (0, 1), 2),
          ("BOTTOMPADDING", (0, 1), (0, 1), 0),
          ("LEFTPADDING", (0, 0), (-1, -1), 0),
          ("RIGHTPADDING", (0, 0), (-1, -1), 0),
      ])
  )

  story.append(locked_block_table)

  doc.build(story)
  buffer.seek(0)
  return buffer.getvalue()


# --- 3. DOCX PARSER ENGINE ---
def extract_matter_data(doc_path):
    doc = docx.Document(doc_path)
    text_lines = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_lines.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = [
                cell.text.strip() for cell in row.cells if cell.text.strip()
            ]
            if row_text:
                text_lines.append(" ".join(row_text))

    full_text = "\n".join(text_lines)

    if "uncontested divorce" in full_text.lower():
        matter_type = "UD"
    elif "contested divorce" in full_text.lower():
        matter_type = "CD"
    elif "annulment" in full_text.lower():
        matter_type = "Annulment"
    elif "variation" in full_text.lower():
        matter_type = "Variation"
    else:
        matter_type = "Others"

    app_match = re.search(
        r"Applicant\s*–\s*([^\n\d]+)", full_text, re.IGNORECASE
    )
    res_match = re.search(
        r"Respondent\s*–\s*([^\n\d]+)", full_text, re.IGNORECASE
    )

    applicant_name = (
        re.split(
            r"\s*[-\s–]\s*upload", app_match.group(1), flags=re.IGNORECASE
        )[0]
        .strip()
        .upper()
        if app_match
        else "NIL"
    )
    respondent_name = (
        re.split(
            r"\s*[-\s–]\s*upload", res_match.group(1), flags=re.IGNORECASE
        )[0]
        .strip()
        .upper()
        if res_match
        else "NIL"
    )

    clean_text = full_text.replace("6224 1848", "").replace("6223 3092", "")
    raw_mobiles = re.findall(
        r"\+?\b(?:65|60|1|44)?[ \-]?[89]\d{3}[ \-]?\d{4}\b|\+?60[ \-]?1\d[ \-]?\d{2,3}[ \-]?\d{4}\b|\b01\d[ \-]?\d{2,3}[ \-]?\d{4}\b",
        clean_text,
    )
    mobiles = [re.sub(r"[\s\-+]", "", mob) for mob in raw_mobiles]
    emails = re.findall(
        r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", clean_text
    )

    def format_phone(clean_mob):
        if clean_mob.startswith("01"):
            clean_mob = f"60{clean_mob[1:]}"
        if clean_mob.startswith(("8", "9")) and len(clean_mob) == 8:
            return f"+65 {clean_mob}"
        else:
            return f"+{clean_mob[:2]} {clean_mob[2:]}"

    app_mob = format_phone(mobiles[0]) if len(mobiles) > 0 else "NIL"
    app_email = emails[0] if len(emails) > 0 else "NIL"

    res_mob = format_phone(mobiles[1]) if len(mobiles) > 1 else "NIL"
    res_email = emails[1] if len(emails) > 1 else "NIL"

    referral = "Google Ad"
    if "referral" in full_text.lower():
        ref_match = re.search(
            r"referral\s*[\s,:\-–]\s*(\w+)", full_text, re.IGNORECASE
        )
        if ref_match and "jav" in ref_match.group(1).lower():
            referral = "Javern"
    elif "jav" in full_text.lower():
        referral = "Javern"

    return {
        "matter_type": matter_type,
        "applicant_name": applicant_name,
        "respondent_name": respondent_name,
        "app_mob": app_mob,
        "app_email": app_email,
        "res_mob": res_mob,
        "res_email": res_email,
        "referral": referral,
    }


def render_pdf_preview(pdf_bytes):
    pdf = pdfium.PdfDocument(pdf_bytes)
    for i, page in enumerate(pdf):
        image = page.render(scale=2).to_pil()
        st.image(image, caption=f"Page {i + 1}", use_container_width=True)


def sync_closed_cartons_from_cloud(db_path):
    """Pulls all closed carton_sync/*.json files from the repo and
    ingests any not already recorded in ingested_cartons."""
    paths = repo_sync.list_carton_sync_files()
    conn = sqlite3.connect(db_path, timeout=10)
    cur = conn.cursor()
    already = {row[0] for row in cur.execute("SELECT carton_no FROM ingested_cartons")}
    results = []
    for path in paths:
        data = repo_sync.fetch_carton_file(path)
        if not data or data.get("status") != "closed":
            continue
        c_no = data.get("carton_no", "UNKNOWN")
        if c_no in already:
            continue
        inserted = 0
        for item in data.get("queue", []):
            try:
                cur.execute(
                    """INSERT INTO archive_records
                       (carton_no, file_no, client_name, matter_type, target_metadata, source_file)
                       VALUES (?, ?, ?, ?, ?, ?)""",
                    (c_no, item.get("line1", "-"), item.get("line2", "-"),
                     item.get("line3", "-"), item.get("line4", "-"), path),
                )
                inserted += 1
            except sqlite3.IntegrityError:
                pass
        cur.execute(
            "INSERT OR REPLACE INTO ingested_cartons (carton_no, ingested_at) VALUES (?, ?)",
            (c_no, datetime.now().isoformat()),
        )
        conn.commit()
        if inserted:
            results.append((c_no, inserted))
    conn.close()
    return results


# --- INITIALIZE SESSION STATE ---
if "uploader_key" not in st.session_state:
    st.session_state["uploader_key"] = 0
if "processed_data" not in st.session_state:
    st.session_state["processed_data"] = {}
if "pdf_binary_store" not in st.session_state:
    st.session_state["pdf_binary_store"] = {}
if "synced_records" not in st.session_state:
    st.session_state["synced_records"] = set()

# --- SIDEBAR CONFIGURATION ---
with st.sidebar:
    st.markdown("### ⚙️ SYSTEM SETTINGS")
    with st.expander("🌐 MASTER SHEET SETUP", expanded=False):
        st.markdown("""
            **Connecting a New Master Sheet:**
            1. Open your target Google Sheet.
            2. Share the file with **Editor** permissions to:
        """)
        try:
            if os.path.exists("credentials.json"):
                with open("credentials.json") as f:
                    c_data = json.load(f)
                    svc_email = c_data.get(
                        "client_email",
                        "your-service-account@iam.gserviceaccount.com",
                    )
            else:
                encoded_str = st.secrets["encoded_creds"]
                decoded_bytes = base64.b64decode(encoded_str)
                creds_dict = json.loads(decoded_bytes)
                svc_email = creds_dict.get(
                    "client_email",
                    "your-service-account@iam.gserviceaccount.com",
                )
        except Exception:
            svc_email = "your-service-account@iam.gserviceaccount.com"

        st.code(svc_email, language="text")
        st.caption("Ensure 'Notify people' is unchecked before sharing.")

# --- DUAL-WING TO TRIPLE-WING SUITE TABS ---
tab_intake, tab_crown, tab_locator = st.tabs([
    "📥 Intake Pipeline & Cover Generator",
    "📦 Crown Box Archival Terminal",
    "🔍 Warehouse Archive Locator",
])

# ==============================================================================
# WING 1: INTAKE PIPELINE & COVER GENERATOR
# ==============================================================================
with tab_intake:
    st.markdown(
        "Drop in your completed **Open File Sheet (`.docx`)** to auto-index against the master matrix and generate the vector cover sheet."
    )

    uploaded_file = st.file_uploader(
        "Drop Open File Sheet (.docx) here",
        type=["docx"],
        key=f"uploader_{st.session_state['uploader_key']}",
    )

    if uploaded_file:
        file_id = uploaded_file.name

        if file_id not in st.session_state["processed_data"]:
            with st.spinner(
                "Parsing intake parameters and synchronizing sequence..."
            ):
                extracted = extract_matter_data(uploaded_file)
                sheet = get_google_sheet()
                next_no, target_row, next_idx = get_next_matter_number(sheet)
                today_str = (
                    datetime.now().strftime("%d %B %Y").lstrip("0").upper()
                )

                clients_str = f"APPLICANT - {extracted['applicant_name']}\nRESPONDENT - {extracted['respondent_name']}"
                contacts_str = f"{extracted['app_mob']} {extracted['app_email']}\n{extracted['res_mob']} {extracted['res_email']}".strip()

                initial_pdf = generate_perfect_pdf(
                    next_no,
                    clients_str,
                    contacts_str,
                    extracted["matter_type"],
                    today_str,
                )

                st.session_state["processed_data"][file_id] = {
                    "matter_no": next_no,
                    "target_row": target_row,
                    "next_idx": next_idx,
                    "today_str": today_str,
                    "matter_type": extracted["matter_type"],
                    "applicant_name": extracted["applicant_name"],
                    "respondent_name": extracted["respondent_name"],
                    "app_mob": extracted["app_mob"],
                    "app_email": extracted["app_email"],
                    "res_mob": extracted["res_mob"],
                    "res_email": extracted["res_email"],
                    "referral": extracted["referral"],
                }
                st.session_state["pdf_binary_store"][file_id] = initial_pdf

        active_data = st.session_state["processed_data"][file_id]

        st.markdown("---")
        st.subheader("📄 Generated Vector Cover Preview")
        render_pdf_preview(st.session_state["pdf_binary_store"][file_id])

        st.write("")
        c_approve, c_cancel = st.columns([1, 1])

        with c_approve:
            is_synced = file_id in st.session_state["synced_records"]
            if st.button(
                "✅ Approve & Sync to Google Sheets",
                use_container_width=True,
                type="primary",
                disabled=is_synced,
            ):
                with st.spinner("Logging to Google Sheet Master Matrix..."):
                    sheet = get_google_sheet()
                    cls = f"APPLICANT - {active_data['applicant_name']}\nRESPONDENT - {active_data['respondent_name']}"
                    cnt = f"{active_data['app_mob']} {active_data['app_email']}\n{active_data['res_mob']} {active_data['res_email']}".strip()

                    new_row = [
                        active_data["next_idx"],
                        active_data["today_str"],
                        active_data["matter_no"],
                        active_data["matter_type"],
                        cls,
                        cnt,
                        active_data["referral"],
                        "Yes",
                        "",
                    ]
                    t_row = active_data["target_row"]
                    sheet.update(
                        range_name=f"A{t_row}:I{t_row}", values=[new_row]
                    )

                    st.session_state["synced_records"].add(file_id)
                    st.toast(
                        f"Matrix Synchronized: File {active_data['matter_no']}",
                        icon="🔹",
                    )
                    st.rerun()

        with c_cancel:
            if st.button("❌ Clear Session", use_container_width=True):
                st.session_state["uploader_key"] += 1
                st.session_state["processed_data"] = {}
                st.session_state["pdf_binary_store"] = {}
                st.session_state["synced_records"] = set()
                st.rerun()

        with st.expander("✏️ Modify Matter Parameters", expanded=False):
            with st.form("edit_intake_data_form"):
                col1, col2 = st.columns(2)

                with col1:
                    edit_matter_no = st.text_input(
                        "File / Matter No.", value=active_data["matter_no"]
                    )
                    edit_date = st.text_input(
                        "Opening Date", value=active_data["today_str"]
                    )
                    edit_matter_type = st.selectbox(
                        "Matter Classification",
                        options=[
                            "UD",
                            "CD",
                            "Annulment",
                            "Variation",
                            "Others",
                        ],
                        index=[
                            "UD",
                            "CD",
                            "Annulment",
                            "Variation",
                            "Others",
                        ].index(active_data["matter_type"])
                        if active_data["matter_type"]
                        in ["UD", "CD", "Annulment", "Variation", "Others"]
                        else 0,
                    )
                    edit_referral = st.text_input(
                        "Referral Source", value=active_data["referral"]
                    )

                with col2:
                    edit_app_name = st.text_input(
                        "Applicant Name", value=active_data["applicant_name"]
                    )
                    edit_app_mob = st.text_input(
                        "Applicant Contact Number",
                        value=active_data["app_mob"],
                    )
                    edit_app_email = st.text_input(
                        "Applicant Email", value=active_data["app_email"]
                    )
                    edit_res_name = st.text_input(
                        "Respondent Name", value=active_data["respondent_name"]
                    )
                    edit_res_mob = st.text_input(
                        "Respondent Contact Number",
                        value=active_data["res_mob"],
                    )
                    edit_res_email = st.text_input(
                        "Respondent Email", value=active_data["res_email"]
                    )

                save_submit = st.form_submit_button(
                    "💾 Save & Update Parameters", use_container_width=True
                )

            if save_submit:
                active_data.update({
                    "matter_no": edit_matter_no,
                    "today_str": edit_date,
                    "matter_type": edit_matter_type,
                    "referral": edit_referral,
                    "applicant_name": edit_app_name,
                    "app_mob": edit_app_mob,
                    "app_email": edit_app_email,
                    "respondent_name": edit_res_name,
                    "res_mob": edit_res_mob,
                    "res_email": edit_res_email,
                })

                cls = f"APPLICANT - {edit_app_name}\nRESPONDENT - {edit_res_name}"
                cnt = f"{edit_app_mob} {edit_app_email}\n{edit_res_mob} {edit_res_email}".strip()

                new_pdf_bytes = generate_perfect_pdf(
                    edit_matter_no, cls, cnt, edit_matter_type, edit_date
                )
                st.session_state["pdf_binary_store"][file_id] = new_pdf_bytes

                with st.spinner("Synchronizing revisions to Google Sheets..."):
                    sheet = get_google_sheet()
                    new_row = [
                        active_data["next_idx"],
                        edit_date,
                        edit_matter_no,
                        edit_matter_type,
                        cls,
                        cnt,
                        edit_referral,
                        "Yes",
                        "",
                    ]
                    t_row = active_data["target_row"]
                    sheet.update(
                        range_name=f"A{t_row}:I{t_row}", values=[new_row]
                    )
                    st.session_state["synced_records"].add(file_id)

                st.toast("Cover sheet regenerated & synced!", icon="✅")
                st.rerun()

        if file_id in st.session_state["synced_records"]:
            st.markdown("---")
            m_tag = active_data["matter_no"].replace("/", "_").strip()
            st.download_button(
                label=f"🖨️ Download Print-Ready Cover Sheet (Matter: {m_tag})",
                data=st.session_state["pdf_binary_store"][file_id],
                file_name=f"21Chambers_Cover_{m_tag}.pdf",
                mime="application/pdf",
                use_container_width=True,
                type="primary",
            )

# ==============================================================================
# WING 2: CROWN BOX ARCHIVAL TERMINAL (Frontend Direct Sync Bridge)
# ==============================================================================
with tab_crown:
    st.markdown("### 📦 Crown Box Archival Terminal (Cloud Synced)")

    with open("scanner_ui.html", "r", encoding="utf-8") as f:
        html_content = f.read()

    # Inject secrets securely from Streamlit to the HTML frontend template
    gh_token = st.secrets.get("github_token", "")
    gh_repo = st.secrets.get("github_repo", "zulfikrycheong/21C-Automation")

    crown_scanner_component = (
        html_content.replace("__GH_TOKEN__", gh_token)
                    .replace("__GH_REPO__", gh_repo)
    )

    components.html(crown_scanner_component, height=1050, scrolling=True)

# ==============================================================================
# WING 3: CROWN ARCHIVAL LOCATOR & SEARCH ENGINE
# ==============================================================================
with tab_locator:
    st.markdown("### 🔍 Crown Worldwide Warehouse Archive Locator")
    st.caption(
        "Instant multi-index search across 8,000+ historical firm files, physical cartons, and client jackets."
    )

    db_path = "crown_base.db"

    if not os.path.exists(db_path):
        st.error(
            "⚠️ `crown_base.db` not found in repository root. Place your compiled database file in the project folder."
        )
    else:
        conn = sqlite3.connect(db_path, timeout=10)
        conn.execute("""CREATE TABLE IF NOT EXISTS ingested_cartons (
            carton_no TEXT PRIMARY KEY,
            ingested_at TEXT
        )""")

        # Remove pre-existing duplicate (carton_no, file_no) rows, keeping the earliest one
        conn.execute("""
            DELETE FROM archive_records
            WHERE id NOT IN (
                SELECT MIN(id)
                FROM archive_records
                GROUP BY carton_no, file_no
            )
        """)
        conn.commit()

        conn.execute("CREATE UNIQUE INDEX IF NOT EXISTS idx_unique_carton_file ON archive_records(carton_no, file_no)")
        conn.commit()
        conn.close()

        last_check = st.session_state.get("last_cloud_sync_check")
        now_ts = datetime.now().timestamp()
        if last_check is None or (now_ts - last_check) > 120:
            newly = sync_closed_cartons_from_cloud(db_path)
            st.session_state["last_cloud_sync_check"] = now_ts
            for c_no, cnt in newly:
                st.toast(f"Auto-indexed {cnt} matters from closed carton {c_no}", icon="📦")

        if st.button("🔄 Force Sync Closed Cartons Now"):
            with st.spinner("Checking cloud for closed cartons..."):
                newly = sync_closed_cartons_from_cloud(db_path)
            st.success(f"Indexed {len(newly)} new carton(s)." if newly else "Nothing new to index.")
            st.rerun()

        # Top Metrics Bar & Database Export Tool (Cloud Safeguard)
        col_metrics, col_download = st.columns([3, 1])

        try:
            conn = sqlite3.connect(db_path, timeout=10)
            total_records = conn.execute(
                "SELECT COUNT(*) FROM archive_records"
            ).fetchone()[0]
            total_cartons = conn.execute(
                "SELECT COUNT(DISTINCT carton_no) FROM archive_records"
            ).fetchone()[0]
            conn.close()

            with col_metrics:
                cm1, cm2, cm3 = st.columns(3)
                with cm1:
                    st.metric("Total Indexed Matters", f"{total_records:,}")
                with cm2:
                    st.metric("Total Physical Cartons", f"{total_cartons:,}")
                with cm3:
                    st.metric("Database Engine", "SQLite B-Tree")

            with col_download:
                st.write("")
                with open(db_path, "rb") as db_file:
                    st.download_button(
                        label="📥 Download Master DB",
                        data=db_file,
                        file_name="crown_base.db",
                        mime="application/x-sqlite3",
                        help="Download the updated SQLite database file to commit back to Git after batch scanning.",
                        use_container_width=True,
                    )
        except Exception as e:
            st.warning(f"Metadata read error: {e}")

        # Ingestion Chute for Newly Exported Cartons
        with st.expander("📥 Ingest Newly Exported Carton (.docx)", expanded=False):
            st.caption("Drop any newly exported Crown Box Record (`.docx`) here to permanently index its files into the search catalog with duplicate protection.")
            new_box_file = st.file_uploader("Upload exported carton .docx", type=["docx"], key="ingest_box_uploader")
            if new_box_file:
                if st.button("Slide Carton into Archive Catalog", type="primary"):
                    try:
                        doc = docx.Document(new_box_file)
                        c_no = "UNKNOWN"
                        if doc.tables and len(doc.tables[0].rows) > 0:
                            header_cell = doc.tables[0].rows[0].cells[0].text
                            c_match = re.search(r"Carton No\.?\s*([\s\S]+?)(?=CARTON DETAILS|$)", header_cell, re.IGNORECASE)
                            c_no = re.sub(r"\s+", " ", c_match.group(1)).strip() if c_match else new_box_file.name.replace(".docx", "")
                        
                        inserted = 0
                        duplicates = 0
                        conn = sqlite3.connect(db_path, timeout=10)
                        cur = conn.cursor()
                        
                        # Ensure unique constraint exists to block duplicates safely
                        cur.execute("CREATE UNIQUE INDEX IF NOT EXISTS idx_unique_carton_file ON archive_records(carton_no, file_no)")

                        for t in doc.tables:
                            for row in t.rows[1:]:
                                if len(row.cells) >= 2:
                                    txt = row.cells[1].text.strip()
                                    lines = [l.strip() for l in txt.split("\n") if l.strip()]
                                    f_no, cl, mt, mt_data = "-", "-", "-", "-"
                                    for line in lines:
                                        if line.startswith("1."): f_no = re.sub(r"^1\.?\s*", "", line).strip()
                                        elif line.startswith("2."): cl = re.sub(r"^2\.?\s*", "", line).strip()
                                        elif line.startswith("3."): mt = re.sub(r"^3\.?\s*", "", line).strip()
                                        elif line.startswith("4."): mt_data = re.sub(r"^4\.?\s*", "", line).strip()
                                    if f_no != "-" or cl != "-":
                                        try:
                                            cur.execute(
                                                """INSERT INTO archive_records (carton_no, file_no, client_name, matter_type, target_metadata, source_file)
                                                   VALUES (?, ?, ?, ?, ?, ?)""",
                                                (c_no, f_no, cl, mt, mt_data, new_box_file.name)
                                            )
                                            inserted += 1
                                        except sqlite3.IntegrityError:
                                            duplicates += 1

                        conn.commit()
                        conn.close()
                        st.success(f"Successfully indexed {inserted} matters from {c_no}! ({duplicates} duplicates skipped)")
                        st.rerun()
                    except Exception as err:
                        st.error(f"Ingestion failed: {err}")

        st.markdown("---")

        # Search Bar & Filter Controls
        search_col, filter_col = st.columns([3, 1])
        with search_col:
            search_query = st.text_input(
                "Search Archives",
                placeholder="Search by Client Name, File No (e.g. 201900782), Property Address, or Postal Code...",
                key="archive_search_input",
            )
        with filter_col:
            field_filter = st.selectbox(
                "Filter Field",
                ["All Fields", "Client Name", "File Number", "Address / Case No"],
            )

        if search_query.strip():
            q = f"%{search_query.strip()}%"
            conn = sqlite3.connect(db_path, timeout=10)

            if field_filter == "Client Name":
                sql = """SELECT carton_no, file_no, client_name, matter_type, target_metadata, source_file 
                         FROM archive_records WHERE client_name LIKE ? ORDER BY id DESC LIMIT 100"""
                params = (q,)
            elif field_filter == "File Number":
                sql = """SELECT carton_no, file_no, client_name, matter_type, target_metadata, source_file 
                         FROM archive_records WHERE file_no LIKE ? ORDER BY id DESC LIMIT 100"""
                params = (q,)
            elif field_filter == "Address / Case No":
                sql = """SELECT carton_no, file_no, client_name, matter_type, target_metadata, source_file 
                         FROM archive_records WHERE target_metadata LIKE ? ORDER BY id DESC LIMIT 100"""
                params = (q,)
            else:
                sql = """SELECT carton_no, file_no, client_name, matter_type, target_metadata, source_file 
                         FROM archive_records 
                         WHERE client_name LIKE ? OR file_no LIKE ? OR target_metadata LIKE ? OR carton_no LIKE ?
                         ORDER BY id DESC LIMIT 100"""
                params = (q, q, q, q)

            results = conn.execute(sql, params).fetchall()
            conn.close()

            if results:
                st.success(f"Found **{len(results)}** matching record(s):")

                for r in results:
                    carton, f_no, client, matter, meta, src = r

                    with st.container():
                        st.markdown(
                            f"""
                            <div style="background:#ffffff; border:1px solid #e2e8f0; border-left:4px solid #1e40af; border-radius:8px; padding:12px 16px; margin-bottom:10px; box-shadow:0 1px 2px rgba(0,0,0,0.03);">
                                <div style="display:flex; justify-content:space-between; align-items:center; margin-bottom:6px;">
                                    <span style="font-size:1.05rem; font-weight:700; color:#1e40af;">File: {f_no}</span>
                                    <span style="background:#eff6ff; color:#1e40af; border:1px solid #bfdbfe; font-size:0.8rem; font-weight:700; padding:3px 10px; border-radius:6px;">📦 Carton: {carton}</span>
                                </div>
                                <div style="font-size:0.95rem; font-weight:600; color:#0f172a; margin-bottom:4px;">Client: {client}</div>
                                <div style="font-size:0.85rem; color:#059669; font-weight:600; margin-bottom:4px;">Matter: {matter}</div>
                                {f'<div style="font-size:0.85rem; color:#475569;">📍 Property / Case: {meta}</div>' if meta and meta != '-' else ''}
                                <div style="font-size:0.75rem; color:#94a3b8; margin-top:6px;">Source: {src}</div>
                            </div>
                            """,
                            unsafe_allow_html=True,
                        )
            else:
                st.info(f"No records found matching '{search_query.strip()}'.")
